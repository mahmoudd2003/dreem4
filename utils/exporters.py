# utils/exporters.py
# -*- coding: utf-8 -*-
"""
Export helpers:
- to_docx(text, path): export organized Word (headings + paragraphs)
- to_pdf(text, path): simple PDF export
"""
from typing import Tuple
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import textwrap

def _is_heading(line: str) -> Tuple[bool, int, str]:
    s = line.strip()
    if s.startswith('### '):
        return True, 3, s[4:].strip()
    if s.startswith('## '):
        return True, 2, s[3:].strip()
    if s.startswith('# '):
        return True, 1, s[2:].strip()
    return False, 0, s

def to_docx(text: str, path: str) -> str:
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(12)

    for raw in (text or '').splitlines():
        is_h, level, content = _is_heading(raw)
        if not content:
            continue
        if is_h:
            if level == 1:
                doc.add_heading(content, level=0)
            elif level == 2:
                doc.add_heading(content, level=1)
            else:
                doc.add_heading(content, level=2)
        else:
            doc.add_paragraph(content)

    doc.save(path)
    return path

def to_pdf(text: str, path: str) -> str:
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    margin = 50
    x = margin
    y = height - margin
    max_width = int((width - 2*margin) / 6)  # rough wrapping for 10pt

    for raw in (text or '').splitlines():
        if not raw.strip():
            y -= 14
            if y < margin: c.showPage(); y = height - margin
            continue
        # simple wrap
        for line in textwrap.wrap(raw, width=90):
            c.setFont("Helvetica", 10)
            c.drawString(x, y, line)
            y -= 14
            if y < margin:
                c.showPage()
                y = height - margin

    c.save()
    return path
